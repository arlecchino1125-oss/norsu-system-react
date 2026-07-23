from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(sys.argv[1])

BLACK = RGBColor(0x00, 0x00, 0x00)
NAVY = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x55, 0x55, 0x55)


def set_run_font(run, size=11, color=BLACK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_style_font(style, size, color=BLACK, bold=False):
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def add_paragraph(
    doc,
    text="",
    *,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before=0,
    after=6,
    line_spacing=1.1,
    size=11,
    color=BLACK,
    bold=False,
    keep_with_next=False,
):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    fmt.keep_with_next = keep_with_next
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold)
    return paragraph


doc = Document()
section = doc.sections[0]
section.start_type = WD_SECTION.NEW_PAGE
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

# standard_business_brief preset
normal = doc.styles["Normal"]
set_style_font(normal, 11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

title = doc.styles["Title"]
set_style_font(title, 22, NAVY, True)
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(8)

subtitle = doc.styles["Subtitle"]
set_style_font(subtitle, 11, MUTED)
subtitle.paragraph_format.space_before = Pt(0)
subtitle.paragraph_format.space_after = Pt(12)

heading_1 = doc.styles["Heading 1"]
set_style_font(heading_1, 16, RGBColor(0x2E, 0x74, 0xB5), True)
heading_1.paragraph_format.space_before = Pt(16)
heading_1.paragraph_format.space_after = Pt(8)

heading_2 = doc.styles["Heading 2"]
set_style_font(heading_2, 13, RGBColor(0x2E, 0x74, 0xB5), True)
heading_2.paragraph_format.space_before = Pt(12)
heading_2.paragraph_format.space_after = Pt(6)

heading_3 = doc.styles["Heading 3"]
set_style_font(heading_3, 12, NAVY, True)
heading_3.paragraph_format.space_before = Pt(8)
heading_3.paragraph_format.space_after = Pt(4)

doc.core_properties.title = "Application Letter - Student Welfare Services Office"
doc.core_properties.subject = "Application to NORSU\u2011Guihulngan Campus"
doc.core_properties.author = "Cejie Bustamante"

# proposal_centerpiece-inspired sender block, restrained for a formal application.
add_paragraph(
    doc,
    "Cejie Bustamante",
    align=WD_ALIGN_PARAGRAPH.CENTER,
    after=1,
    line_spacing=1.0,
    size=15,
    color=NAVY,
    bold=True,
    keep_with_next=True,
)
add_paragraph(
    doc,
    "09606874599  |  cjbustajod@gmail.com",
    align=WD_ALIGN_PARAGRAPH.CENTER,
    after=14,
    line_spacing=1.0,
    size=9.5,
    color=MUTED,
)

add_paragraph(doc, "23 July 2026", after=12, line_spacing=1.0, keep_with_next=True)

recipient_lines = [
    ("Hon. Noel Marjon E. Yasi, Psy.D.", True),
    ("University President", False),
    ("Negros Oriental State University", False),
    ("Guihulngan Campus", False),
    ("Guihulngan City, Negros Oriental", False),
]
for index, (text, bold) in enumerate(recipient_lines):
    add_paragraph(
        doc,
        text,
        after=12 if index == len(recipient_lines) - 1 else 0,
        line_spacing=1.0,
        bold=bold,
        keep_with_next=True,
    )

add_paragraph(doc, "Dear Mr. President,", after=10, line_spacing=1.0, keep_with_next=True)

body_paragraphs = [
    (
        "I am writing to apply for a position in the Student Welfare Services Office "
        "(SWSO) at NORSU\u2011Guihulngan Campus. As a recent Bachelor of Science in "
        "Computer Science graduate of NORSU\u2011Guihulngan, I am eager to contribute my "
        "technical skills, experience, and familiarity with student welfare operations "
        "in service to the university and its students."
    ),
    (
        "During my on-the-job training at the CARE Center Office in 2025, I gained "
        "firsthand knowledge of its daily operations, including assisting students, "
        "addressing their concerns, and connecting them with appropriate services. "
        "After observing the challenges involved in managing student records and "
        "requests, I took the initiative to design and develop the CARE Center "
        "Management System specifically for NORSU\u2011Guihulngan Campus. The system was "
        "created to organize student information, manage records, and support the "
        "efficient delivery of the office's services. I am prepared to maintain and "
        "improve the system according to the office's needs."
    ),
    (
        "I also completed an internship as an IT Support Intern at Qualfon, where I "
        "gained hands-on experience in IT inventory management and assisting users "
        "with technical concerns. This experience strengthened my technical, "
        "organizational, communication, and problem-solving skills and taught me to "
        "provide dependable assistance in a professional environment."
    ),
    (
        "My experience has prepared me to assist students, respond respectfully to "
        "their concerns, maintain accurate and confidential records, and support the "
        "office's daily operations. With my knowledge of the CARE Center, experience "
        "developing its management system, and ability to provide practical technical "
        "solutions, I am confident that I can contribute to a more organized, "
        "efficient, and student-centered office."
    ),
    (
        "I would welcome the opportunity to discuss how my qualifications and "
        "experience could contribute to the Student Welfare Services Office. My "
        "resume and supporting documents are enclosed for your consideration."
    ),
    (
        "Thank you for considering my application. I look forward to the opportunity "
        "to serve NORSU\u2011Guihulngan."
    ),
]

for text in body_paragraphs:
    add_paragraph(doc, text)

add_paragraph(doc, "Respectfully yours,", after=24, line_spacing=1.0, keep_with_next=True)
add_paragraph(doc, "Cejie Bustamante", after=0, line_spacing=1.0, bold=True)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
